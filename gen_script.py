import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

doc = Document()

def h(text, level=2):
    doc.add_heading(text, level=level)

def b(text):
    doc.add_paragraph(text, style='List Bullet')

def note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

def sep():
    doc.add_paragraph()

# TITLE
doc.add_heading('VIDEO DEMO SCRIPT — Expense Approval AI System', 0)
note('Band of Agents Hackathon 2026 | Track 1 | ~2 phut | Quay man hinh + noi')
sep()

# INTRO
h('INTRO  (~10 giay)', 2)
b('Mo Employee Dashboard: http://103.157.204.120:5000')
note('Noi: "Day la he thong duyet chi phi doanh nghiep tu dong — 4 AI agent tren Band, dung GPT-4o va Llama 3.1 70B."')
sep()

# SCENE 1
h('Kich ban 1 — Rui ro THAP, tu dong duyet  (~35 giay)', 2)
h('Dien form:', 3)
b('Ten: [ten ban]  |  So tien: $200  |  Phong ban: HR')
b('Nha cung cap: Staples  |  Mo ta: office supplies')
b('Bam Submit')
h('Chuyen qua Band room — show 4 agent chay:', 3)
b('Budget Checker: kiem tra ngan sach HR -> OK')
b('Policy Checker (Llama 3.1 70B): khong vi pham -> COMPLIANT')
b('Risk Evaluator: $200 < $500, COMPLIANT -> LOW RISK -> AUTO-APPROVE')
b('Approval Notifier: gui thong bao APPROVED')
note('Noi: "$200 duoi nguong $500, tuan thu chinh sach — he thong tu duyet ~15 giay, khong can nguoi."')
sep()

# SCENE 2
h('Kich ban 2 — Rui ro TRUNG BINH, manager duyet  (~35 giay)', 2)
h('Dien form:', 3)
b('So tien: $800  |  Phong ban: Engineering')
b('Nha cung cap: Amazon  |  Mo ta: AWS software license')
b('Bam Submit')
h('Band room:', 3)
b('Risk Evaluator: $800 trong khoang $500-$1,500 -> MEDIUM -> can manager')
h('Chuyen Admin Panel: http://103.157.204.120:5001', 3)
b('Tab "Cho Duyet" -> thay expense EXP-XXXXXX day du thong tin')
b('Bam Duyet -> toast xanh, status doi sang APPROVED')
b('Band room: thong bao APPROVED voi ten + so tien that hien ra')
note('Noi: "Manager duyet qua web, khong can tai khoan Band. Cap nhat real-time."')
sep()

# SCENE 3
h('Kich ban 3 — Rui ro CAO, leo thang CFO  (~25 giay)', 2)
h('Dien form:', 3)
b('So tien: $6,000  |  Phong ban: Finance  |  Mo ta: SAP ERP license')
h('Band room:', 3)
b('Policy Checker: >$5,000 -> NON-COMPLIANT (can CFO)')
b('Risk Evaluator: HIGH RISK -> CFO ESCALATION')
h('Show Admin Panel — 2 tab nhanh:', 3)
b('Live Feed: xem toan bo cuoc tro chuyen agent tu audit log')
b('Ngan Sach: thanh tien do tung phong ban, nut Reset ngan sach')
b('Bam Tu choi voi ly do -> Band room nhan thong bao REJECTED')
note('Noi: "$6,000 vi pham policy CFO sign-off — he thong tu leo thang, khong bo sot."')
sep()

# OUTRO
h('OUTRO  (~10 giay)', 2)
note('Noi: "4 agent, 2 AI provider, deploy live 24/7. Demo ket thuc."')
b('Hien thi: http://103.157.204.120:5000')
b('GitHub: github.com/thanhvuaws-jpg/band-expense-approval-agents')
sep()

# TECH NOTES
h('Tech Stack (cho judge)', 2)
b('Band SDK: LangGraphAdapter + InMemorySaver')
b('GPT-4o via AIML API — Budget Checker, Risk Evaluator, Approval Notifier')
b('Llama 3.1 70B via Featherless AI — Policy Checker')
b('Admin Panel: direct DB update + thong bao ra Band room (khong phu thuoc Approval Notifier)')
b('Live Feed: doc tu audit_log SQLite (khong phai Band REST API)')
b('Docker Compose — 3 container: agents, dashboard, admin')

doc.save('Demo_Script_Video.docx')
print('OK')
